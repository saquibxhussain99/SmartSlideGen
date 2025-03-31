# figure_extractor.py

import os
import re
import uuid
from docx import Document
from lxml import etree  # You need lxml if you want to climb the XML tree


# def extract_figures_from_docx(docx_path, output_folder="extracted_figures"):
#     """
#     Extract images from inline shapes in a .docx, saving them into `output_folder`.
#     Returns metadata about each extracted image.
#     """
#     doc = Document(docx_path)
#     os.makedirs(output_folder, exist_ok=True)

#     figures_info = []
#     figure_counter = 1

#     # python-docx gives access to images via inline_shapes
#     for inline_shape in doc.inline_shapes:
#         # Each inline shape has an image reference stored in _inline.graphic.graphicData...
#         # We can get the relationship ID (rId) which points to the actual image data
#         rId = inline_shape._inline.graphic.graphicData.pic.blipFill.blip.embed

#         # Retrieve the image part from the doc's relationships
#         image_part = doc.part.related_parts[rId]
#         image_bytes = image_part.blob

#         # Build a unique filename
#         image_filename = f"figure_{figure_counter}_{uuid.uuid4().hex[:8]}.png"
#         image_path = os.path.join(output_folder, image_filename)
#         with open(image_path, 'wb') as f:
#             f.write(image_bytes)

#         # If you want to find text around the image, you can try to reference
#         # inline_shape._inline to see the parent or paragraph. 
#         # But it can be tricky—there’s no direct "paragraph text" property for inline shapes.
#         # Typically you might gather that from your doc.paragraphs separately.

#         figures_info.append({
#             'id': figure_counter,
#             'type': 'image',
#             'path': image_path,
#             # 'paragraph_text': ...
#         })
#         figure_counter += 1

#     # You can add additional logic for tables, shapes, or other embedded objects here.

#     return figures_info


# def extract_figures_from_docx(docx_path, output_folder="extracted_figures"):
#     """
#     Extract images (inline shapes) from a .docx, saving them into `output_folder`.
#     Returns a list of figure metadata, including 'paragraph_text'.
#     """
#     doc = Document(docx_path)
#     os.makedirs(output_folder, exist_ok=True)

#     figures_info = []
#     figure_counter = 1

#     # Iterate over each paragraph so we can capture the paragraph text
#     for paragraph in doc.paragraphs:
#         paragraph_text = paragraph.text.strip()
        
#         # For each inline shape (image) in this paragraph, extract and save it
#         for inline_shape in paragraph.inline_shapes:
#             # Relationship ID for the embedded image
#             rId = inline_shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            
#             # Retrieve the image part from the document relationships
#             image_part = doc.part.related_parts[rId]
#             image_bytes = image_part.blob

#             # Build a unique filename
#             image_filename = f"figure_{figure_counter}_{uuid.uuid4().hex[:8]}.png"
#             image_path = os.path.join(output_folder, image_filename)
            
#             # Write the image bytes to disk
#             with open(image_path, 'wb') as f:
#                 f.write(image_bytes)

#             # Store metadata, including the paragraph text
#             figures_info.append({
#                 'id': figure_counter,
#                 'type': 'image',
#                 'path': image_path,
#                 'paragraph_text': paragraph_text
#             })
#             figure_counter += 1

#     # (Optional) Add logic for tables, shapes, or anchored shapes here

#     return figures_info


def extract_figures_from_docx(docx_path, output_folder="extracted_figures"):
    doc = Document(docx_path)
    os.makedirs(output_folder, exist_ok=True)

    figures_info = []
    figure_counter = 1

    for shape in doc.inline_shapes:
        # 1) Extract the image
        rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        image_part = doc.part.related_parts[rId]
        image_bytes = image_part.blob

        image_filename = f"figure_{figure_counter}_{uuid.uuid4().hex[:8]}.png"
        image_path = os.path.join(output_folder, image_filename)
        with open(image_path, 'wb') as f:
            f.write(image_bytes)

        # 2) Find the paragraph element that contains this shape
        # shape._inline is an oxml element; we climb the tree until we find a w:p node
        p_xml_element = shape._inline
        paragraph_text = ""

        while p_xml_element is not None:
            if p_xml_element.tag.endswith('p'):
                # We found the paragraph element
                # Gather all text from within this paragraph
                paragraph_text = "".join(
                    node.text 
                    for node in p_xml_element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                    if node.text
                ).strip()
                break
            p_xml_element = p_xml_element.getparent()

        # 3) Store results
        figures_info.append({
            "id": figure_counter,
            "type": "image",
            "path": image_path,
            "paragraph_text": paragraph_text
        })
        figure_counter += 1

    return figures_info



def decide_slide_mapping(figures_info, text_slides):
    """
    Given a list of figure metadata and your text-based slides (or paragraphs),
    decide which slide each figure should go on.

    :param figures_info: List of dicts describing figures (id, type, path, paragraph_text, etc.)
    :param text_slides: List of text slides. Each slide is a dict or object with keys like {'title':..., 'content':...}
    :return: A mapping or updated list that shows which slide each figure belongs to.
    """

    # For demonstration, we do a naive approach:
    # 1) We look at paragraph_text for keywords to match slide content
    # 2) If figure is "large" or "detailed", we might put it on its own slide
    # 3) Otherwise, we attach it to the best matching slide

    figure_to_slide_map = []

    for fig in figures_info:
        # Decide if "large" or "detailed" => For now, let's just say "table" or something is large
        # In a real scenario, you might measure the image’s dimensions or parse table size
        is_large = (fig['type'] == 'table')

        # Attempt a naive text match to find a suitable slide
        best_slide_index = None
        best_match_score = 0

        fig_text = fig['paragraph_text'] if fig['type'] == 'image' else ""

        for i, slide in enumerate(text_slides):
            slide_content = (slide.get('title', "") + " " + slide.get('content', "")).lower()
            # Simple example: count how many words overlap
            overlap = 0
            for word in re.split(r'\W+', fig_text.lower()):
                if word in slide_content:
                    overlap += 1

            if overlap > best_match_score:
                best_match_score = overlap
                best_slide_index = i

        # If we found a good match and the figure is not "large", place it on that slide
        # If large or no match found, create a new slide
        if best_slide_index is not None and not is_large:
            figure_to_slide_map.append({
                'figure_id': fig['id'],
                'figure_type': fig['type'],
                'figure_path': fig['path'] if fig['type'] == 'image' else None,
                'slide_index': best_slide_index,
                'own_slide': False
            })
        else:
            # Make a new "slide" for this figure
            # In an actual pipeline, you might append to text_slides, or store them separately
            figure_to_slide_map.append({
                'figure_id': fig['id'],
                'figure_type': fig['type'],
                'figure_path': fig['path'] if fig['type'] == 'image' else None,
                'slide_index': None,  # or len(text_slides) if you create a new slide
                'own_slide': True
            })

    return figure_to_slide_map


if __name__ == "__main__":
    """
    Example usage:
    
    1. Extract figures from a sample docx
    2. Suppose we have some text slides from your existing text extraction step
    3. Get a mapping for where each figure should go
    """
    sample_docx = "example.docx"  # Replace with your .docx file
    figures = extract_figures_from_docx(sample_docx, output_folder="temp_figures")

    # Example set of slides from your text-extraction logic
    text_slides_example = [
        {'title': 'Introduction', 'content': 'This presentation covers the project overview and objectives.'},
        {'title': 'Methodology', 'content': 'We used a step-by-step process to analyze the data and produce results.'},
        {'title': 'Results', 'content': 'The final results are shown in the chart below.'}
    ]

    mapping = decide_slide_mapping(figures, text_slides_example)

    print("Extracted Figures:")
    for f in figures:
        print(f"  {f}")

    print("\nFigure-to-Slide Mapping:")
    for m in mapping:
        print(m)
